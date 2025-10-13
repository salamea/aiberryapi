"""
Document Processing Service
Handles document upload, parsing, chunking, and storage
"""
import logging
import uuid
from typing import List, Dict, Any
from io import BytesIO

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import PyPDF2
from docx import Document as DocxDocument

from ..config import Settings
from .vector_service import VectorService

logger = logging.getLogger(__name__)


class DocumentService:
    """Service for document processing and management"""

    def __init__(self, settings: Settings, vector_service: VectorService):
        self.settings = settings
        self.vector_service = vector_service

        # Text splitter for chunking
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )

    async def process_document(
        self,
        filename: str,
        content: bytes,
        file_type: str
    ) -> Dict[str, Any]:
        """
        Process uploaded document

        Args:
            filename: Original filename
            content: File content as bytes
            file_type: File extension (.pdf, .docx, .txt)

        Returns:
            Dict with processing results
        """
        try:
            # Validate file size
            file_size_mb = len(content) / (1024 * 1024)
            if file_size_mb > self.settings.max_file_size_mb:
                raise ValueError(f"File size exceeds maximum of {self.settings.max_file_size_mb}MB")

            # Extract text based on file type
            if file_type == '.pdf':
                text = await self._extract_pdf_text(content)
            elif file_type == '.docx':
                text = await self._extract_docx_text(content)
            elif file_type == '.txt':
                text = content.decode('utf-8')
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

            # Validate extracted text
            if not text.strip():
                raise ValueError("No text could be extracted from the document")

            # Create document ID
            document_id = str(uuid.uuid4())

            # Split into chunks
            chunks = self.text_splitter.split_text(text)

            # Create Document objects
            documents = [
                Document(
                    page_content=chunk,
                    metadata={
                        'filename': filename,
                        'document_id': document_id,
                        'chunk_index': idx,
                        'total_chunks': len(chunks),
                        'file_type': file_type
                    }
                )
                for idx, chunk in enumerate(chunks)
            ]

            # Store in vector database
            chunks_created = await self.vector_service.add_documents(
                documents=documents,
                document_id=document_id
            )

            logger.info(f"Processed document {filename}: {chunks_created} chunks created")

            return {
                'document_id': document_id,
                'filename': filename,
                'chunks_created': chunks_created,
                'status': 'success'
            }

        except Exception as e:
            logger.error(f"Error processing document: {e}")
            raise

    async def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF"""
        try:
            pdf_file = BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text_parts = []
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())

            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    async def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            docx_file = BytesIO(content)
            doc = DocxDocument(docx_file)

            text_parts = [paragraph.text for paragraph in doc.paragraphs]
            return "\n\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    async def delete_document(self, document_id: str) -> bool:
        """Delete document from vector store"""
        try:
            return await self.vector_service.delete_document(document_id)
        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            return False

    async def list_documents(self) -> List[Dict[str, Any]]:
        """List all documents"""
        try:
            return await self.vector_service.list_documents()
        except Exception as e:
            logger.error(f"Error listing documents: {e}")
            return []

    async def search_documents(
        self,
        query: str,
        k: int = 5
    ) -> List[Dict[str, Any]]:
        """Search documents by query"""
        try:
            results = await self.vector_service.similarity_search(
                query=query,
                k=k,
                score_threshold=self.settings.similarity_threshold
            )

            return [
                {
                    'content': doc.page_content,
                    'filename': doc.metadata.get('filename'),
                    'document_id': doc.metadata.get('document_id'),
                    'score': doc.metadata.get('score')
                }
                for doc in results
            ]

        except Exception as e:
            logger.error(f"Error searching documents: {e}")
            return []
